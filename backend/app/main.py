"""
FastAPI application: the API layer for Ask My Docs.

Endpoints:
  GET  /health                    - liveness check
  GET  /metrics                    - Prometheus metrics (scraped by Grafana stack)
  GET  /api/documents              - list the patient's documents
  GET  /api/documents/{doc_id}     - full text of one document
  POST /api/documents/upload       - upload one or more text documents
  POST /api/ask                    - ask a question, get a cited answer
  POST /api/feedback               - send +1 / -1 feedback on an answer

Also mounts the static frontend (frontend/) at "/", so a single process
serves both the API and the UI -- run with `python main.py` from the repo
root.
"""

import logging
import re
from contextlib import asynccontextmanager
from io import BytesIO
from pathlib import Path

from fastapi import FastAPI, File, HTTPException, Response, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles

from app.config import FRONTEND_DIR
from app import db, metrics
from app.ingest import load_document_from_text, new_id
from app.models import (
    AskRequest,
    AskResponse,
    Citation,
    DocumentDetail,
    DocumentSummary,
    FeedbackRequest,
    FeedbackResponse,
    UploadDocumentsResponse,
    UploadIssue,
)
from app.rag import RagPipeline

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("ask_my_docs")
SUPPORTED_UPLOAD_EXTENSIONS = {".txt", ".md", ".text", ".pdf", ".docx"}

# Built once at import time; shared by every request and by eval scripts.
pipeline = RagPipeline()
db.init_db()


@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info(
        "Loaded %d documents / %d chunks. vector_backend=%s rerank_backend=%s",
        len(pipeline.documents),
        len(pipeline.chunks),
        pipeline.retriever.vector.backend_name,
        pipeline.reranker.backend_name,
    )
    logger.info("Prometheus metrics available at /metrics")
    yield


app = FastAPI(
    title="Ask My Docs API",
    description="Healthcare RAG API: hybrid retrieval + reranking + citation "
    "enforcement over a patient's own medical documents.",
    version="0.1.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
def health():
    return {
        "status": "ok",
        "documents_loaded": len(pipeline.documents),
        "chunks_indexed": len(pipeline.chunks),
        "vector_backend": pipeline.retriever.vector.backend_name,
        "rerank_backend": pipeline.reranker.backend_name,
    }


@app.get("/metrics")
def prometheus_metrics():
    """Scraped by the Prometheus service in docker-compose.yml / monitoring/."""
    return Response(content=metrics.render_metrics(), media_type=metrics.METRICS_CONTENT_TYPE)


@app.get("/api/documents", response_model=list[DocumentSummary])
def list_documents():
    return [
        DocumentSummary(doc_id=d.doc_id, title=d.title, doc_type=d.doc_type, date=d.date)
        for d in pipeline.list_documents()
    ]


@app.get("/api/documents/{doc_id}", response_model=DocumentDetail)
def get_document(doc_id: str):
    doc = pipeline.get_document(doc_id)
    if not doc:
        metrics.record_error("get_document")
        raise HTTPException(status_code=404, detail="Document not found")
    return DocumentDetail(
        doc_id=doc.doc_id,
        title=doc.title,
        doc_type=doc.doc_type,
        date=doc.date,
        content=doc.content,
    )


def _safe_uploaded_doc_id(filename: str) -> str:
    stem = Path(filename).stem or "uploaded_document"
    normalized = re.sub(r"[^a-zA-Z0-9_-]+", "_", stem.strip().lower()).strip("_")
    return f"upload_{normalized or 'document'}_{new_id()[:8]}"


def _extract_uploaded_text(filename: str, raw: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext in {".txt", ".md", ".text"}:
        return raw.decode("utf-8")

    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(raw))
        pages = []
        for page in reader.pages:
            pages.append((page.extract_text() or "").strip())
        return "\n\n".join(p for p in pages if p)

    if ext == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(BytesIO(raw))
        lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)

    raise ValueError("Unsupported file type")


@app.post("/api/documents/upload", response_model=UploadDocumentsResponse)
async def upload_documents(files: list[UploadFile] = File(...)):
    uploaded_docs = []
    skipped = []

    for f in files:
        name = f.filename or "uploaded_document"
        ext = Path(name).suffix.lower()
        if ext not in SUPPORTED_UPLOAD_EXTENSIONS:
            skipped.append(
                UploadIssue(
                    filename=name,
                    reason="Unsupported file type. Upload .txt, .md, .text, .pdf, or .docx files.",
                )
            )
            continue

        raw = await f.read()
        if not raw:
            skipped.append(UploadIssue(filename=name, reason="File is empty."))
            continue

        try:
            text = _extract_uploaded_text(name, raw)
        except UnicodeDecodeError:
            skipped.append(
                UploadIssue(
                    filename=name,
                    reason="Couldn't read this file as UTF-8 text.",
                )
            )
            continue
        except Exception:
            skipped.append(
                UploadIssue(
                    filename=name,
                    reason="Couldn't extract text from this file.",
                )
            )
            continue

        title = Path(name).stem.replace("_", " ").replace("-", " ").strip() or name
        doc = load_document_from_text(
            doc_id=_safe_uploaded_doc_id(name),
            raw=text,
            title=title,
            doc_type="Uploaded Document",
        )
        if not doc.chunks:
            skipped.append(
                UploadIssue(
                    filename=name,
                    reason="No readable text content found.",
                )
            )
            continue
        uploaded_docs.append(doc)

    pipeline.add_documents(uploaded_docs)

    return UploadDocumentsResponse(
        uploaded_count=len(uploaded_docs),
        uploaded=[
            DocumentSummary(
                doc_id=d.doc_id,
                title=d.title,
                doc_type=d.doc_type,
                date=d.date,
            )
            for d in uploaded_docs
        ],
        skipped=skipped,
    )


@app.post("/api/ask", response_model=AskResponse)
def ask(req: AskRequest):
    if req.document_id and not pipeline.get_document(req.document_id):
        metrics.record_error("ask")
        raise HTTPException(status_code=404, detail="Document not found")

    try:
        result = pipeline.answer(req.question, document_id=req.document_id)
    except Exception:
        metrics.record_error("ask")
        raise

    conversation_id = req.conversation_id or new_id()

    db.save_conversation(
        conversation_id=conversation_id,
        question=req.question,
        answer=result["answer"],
        document_id=req.document_id,
        grounded=result["grounded"],
        citations=result["citations"],
        retrieval_debug=result["retrieval_debug"],
    )

    return AskResponse(
        conversation_id=conversation_id,
        question=req.question,
        answer=result["answer"],
        citations=[Citation(**c) for c in result["citations"]],
        disclaimer=result["disclaimer"],
        grounded=result["grounded"],
        retrieval_debug=result["retrieval_debug"],
    )


@app.post("/api/feedback", response_model=FeedbackResponse)
def feedback(req: FeedbackRequest):
    db.save_feedback(req.conversation_id, req.feedback, req.comment)
    metrics.record_feedback(req.feedback)
    return FeedbackResponse(
        message=f"Feedback received for conversation {req.conversation_id}: {req.feedback}"
    )


# --- Static frontend ------------------------------------------------------
# Mounted last so it doesn't shadow the /api routes above.
if FRONTEND_DIR.exists():
    app.mount("/", StaticFiles(directory=str(FRONTEND_DIR), html=True), name="frontend")
